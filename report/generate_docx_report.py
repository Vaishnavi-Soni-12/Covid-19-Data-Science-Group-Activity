from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_run_style(run, font_name='Times New Roman', font_size=12, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def add_paragraph(doc, text, justify=True, bold=False, size=12):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        set_run_style(run, font_size=size, bold=bold)
    return p

def add_code_snippet(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(50, 50, 50)

def add_image_with_explanation(doc, img_name, caption, explanation):
    img_path = f"report/assets/{img_name}"
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        cap_p = doc.add_paragraph(caption)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap_p.runs:
            set_run_style(run, font_size=10, italic=True)
        
        exp_p = doc.add_paragraph(explanation)
        exp_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in exp_p.runs:
            set_run_style(run, font_size=11)
    else:
        print(f"Warning: Image {img_path} not found.")

def create_corporate_report():
    doc = Document()
    
    # --- Cover Page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('\n\n\nCORPORATE ANALYSIS REPORT:\nTHE GLOBAL IMPACT OF COVID-19')
    set_run_style(run, font_size=28, bold=True)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('Comprehensive Data Science Retrospective & Policy Discussion')
    set_run_style(run, font_size=16)
    
    doc.add_paragraph('\n' * 5)
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run('Author: Antigravity AI Systems\nProject: Global Health Intelligence\nDate: May 06, 2026')
    set_run_style(run, font_size=12)
    
    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    add_paragraph(doc, 
        "This report delivers a deep-dive analysis of the COVID-19 pandemic using the OWID global repository. "
        "We have processed 16 high-dimensional visualizations to explain the spread mechanics, demographic vulnerabilities, "
        "and the transformative success of global vaccination efforts. The analysis confirms that while the virus spread "
        "indiscriminately, the impact was determined by local infrastructure, age distribution, and political stringency."
    )

    doc.add_page_break()

    # --- 1. Introduction & Objective ---
    doc.add_heading('1. Introduction', level=1)
    add_paragraph(doc, 
        "The COVID-19 pandemic caused by SARS-CoV-2 represents the most documented epidemiological event in human history. "
        "This study serves as a retrospective analysis to extract actionable insights from multi-year data."
    )
    
    doc.add_heading('2. Objectives', level=1)
    add_paragraph(doc, 
        "The objective is to utilize end-to-end data science—from automated ingestion to visual correlation—to answer "
        "critical questions regarding the pandemic's cause, its global outcome, and the 'suffering index' of different nations."
    )

    doc.add_page_break()

    # --- 3. Spread Analysis & Wave Periodicity ---
    doc.add_heading('3. Results: Global Spread Trends', level=1)
    
    add_image_with_explanation(doc, "global_total_cases.png", "Figure 1: Cumulative Case Growth", 
        "Explanation: This graph shows the exponential rise in COVID-19 cases globally. The curve illustrates the sheer scale of the pandemic, "
        "moving from a few thousand cases in early 2020 to over 700 million by 2024. The relatively linear growth in the later stages "
        "reflects the transition from an acute emergency to an endemic state.")

    add_image_with_explanation(doc, "global_total_deaths.png", "Figure 2: Global Cumulative Mortality", 
        "Explanation: Mortality followed the case growth but with distinct shifts in slope. The sharpest increases correspond to periods "
        "before widespread vaccination and the emergence of more virulent strains like Delta.")

    add_image_with_explanation(doc, "daily_new_cases.png", "Figure 3: Daily New Cases Wave Analysis", 
        "Explanation: This bar chart captures the massive waves of infection. The largest spike seen in early 2022 represents the Omicron wave, "
        "which was characterized by high transmissibility but lower average severity compared to previous variants.")

    add_image_with_explanation(doc, "daily_new_deaths.png", "Figure 4: Daily Mortality Fluctuations", 
        "Explanation: Unlike new cases, daily deaths reached their peaks earlier in the pandemic. This shows that later waves, despite having "
        "more cases, resulted in fewer deaths relative to the volume, indicating improved clinical management and vaccine protection.")

    doc.add_page_break()

    # --- 4. Geographic & Comparative Outcomes ---
    doc.add_heading('4. Comparative Outcomes: Regional Analysis', level=1)

    add_image_with_explanation(doc, "top_10_cases.png", "Figure 5: Highest Case Loads by Nation", 
        "Explanation: The United States, India, and Brazil lead the global rankings. This reflects a combination of large populations, "
        "high international travel hubs, and varying levels of initial containment success.")

    add_image_with_explanation(doc, "top_10_deaths.png", "Figure 6: Highest Mortality by Nation", 
        "Explanation: Total deaths align closely with case volume, but nations like Brazil show a higher mortality burden relative to their case "
        "rankings, suggesting a higher strain on their intensive care infrastructure during peaks.")

    add_image_with_explanation(doc, "cfr_top_countries.png", "Figure 7: Case Fatality Rate (CFR) Comparison", 
        "Explanation: The CFR identifies the 'deadliness' of the pandemic in specific contexts. High CFRs in certain regions indicate "
        "under-reporting of mild cases or a lack of access to oxygen and advanced therapies.")

    add_image_with_explanation(doc, "continent_distribution.png", "Figure 8: Total Deaths by Continent", 
        "Explanation: Europe and the Americas bore the brunt of the reported mortality. This may be influenced by reporting transparency "
        "and the high proportion of elderly citizens in these regions compared to Africa and parts of Asia.")

    doc.add_page_break()

    # --- 5. Socio-Economic & Demographic Drivers ---
    doc.add_heading('5. Socio-Economic & Demographic Analysis', level=1)

    add_image_with_explanation(doc, "gdp_vs_mortality.png", "Figure 9: Economic Wealth vs. Mortality per Million", 
        "Explanation: This scatter plot reveals a surprising trend: higher GDP countries often had higher reported mortality. "
        "Discussion suggests this is due to more accurate cause-of-death reporting and higher international connectivity in wealthy nations.")

    add_image_with_explanation(doc, "elderly_vs_mortality.png", "Figure 10: Demographic Aging vs. Mortality", 
        "Explanation: This is one of the strongest correlations in the dataset. Countries with higher percentages of citizens over 65 "
        "faced drastically higher mortality rates, confirming that age was the single greatest risk factor.")

    add_image_with_explanation(doc, "life_expectancy_correlation.png", "Figure 11: Health Infrastructure vs. Death Rates", 
        "Explanation: Similar to the GDP chart, countries with higher life expectancy (and thus more elderly citizens) were more vulnerable "
        "to high mortality counts during the pandemic peaks.")

    add_image_with_explanation(doc, "testing_vs_positivity.png", "Figure 12: Testing Density impact", 
        "Explanation: High testing density generally led to lower positive rates, as more asymptomatic and mild cases were captured, "
        "allowing for better isolation and control measures.")

    doc.add_page_break()

    # --- 6. Intervention & Efficacy ---
    doc.add_heading('6. Intervention & Vaccine Efficacy', level=1)

    add_image_with_explanation(doc, "vaccination_growth.png", "Figure 13: Global Vaccination Progress", 
        "Explanation: The rapid rise in vaccinations from 2021 onwards represents a historic public health achievement. "
        "This curve correlates with the eventual decoupling of case growth from mortality growth.")

    add_image_with_explanation(doc, "vax_vs_deaths.png", "Figure 14: Vaccine Impact on Mortality", 
        "Explanation: Countries with higher 'People Fully Vaccinated per Hundred' consistently show lower mortality rates in the latter "
        "stages of the pandemic, proving the strategic value of the vaccine rollout.")

    add_image_with_explanation(doc, "stringency_vs_cases.png", "Figure 15: Policy Stringency vs. Case Load", 
        "Explanation: Using India as a case study, we see how high stringency (lockdowns/restrictions) was used to blunt the force "
        "of initial waves. The easing of restrictions often preceded new surges.")

    add_image_with_explanation(doc, "hospitalization_trends.png", "Figure 16: ICU and Hospitalization Burden", 
        "Explanation: This graph tracks the direct 'suffering' in the healthcare system. The peaks in hospital patients show where "
        "the risk of system collapse was highest, necessitating emergency interventions.")

    doc.add_page_break()

    # --- 7. Final Results & Discussion ---
    doc.add_heading('7. Final Results and Discussion', level=1)
    
    add_paragraph(doc, "Major Findings:", bold=True)
    add_paragraph(doc, 
        "1. The pandemic transitioned through three phases: The pre-vaccine era (high CFR), the variant surge (Delta/Omicron), "
        "and the endemic phase (high vaccination, low relative mortality).\n"
        "2. Demographic profiling (Age 65+) was the primary predictor of mortality, far outweighing GDP or healthcare spend alone.\n"
        "3. Vaccine efficacy is clearly visible in the data, where high-vax populations saw a significant 'flattening' of the death curve "
        "relative to the volume of new infections.")

    add_paragraph(doc, "Discussion:", bold=True)
    add_paragraph(doc, 
        "The analysis demonstrates that COVID-19's impact was as much a social and political event as a biological one. "
        "Wealthy nations, despite their resources, were vulnerable due to their age distributions and interconnectedness. "
        "The 'suffering' documented in the hospitalization data shows that capacity management was the deciding factor in survival. "
        "Future preparedness must focus on rapid diagnostic scaling (testing) and protecting high-risk demographic clusters identified here.")

    # --- 8. Conclusion ---
    doc.add_heading('8. Conclusion', level=1)
    add_paragraph(doc, 
        "In conclusion, this data science project successfully mapped the end-to-end trajectory of COVID-19. "
        "The results provide a technical validation of public health interventions and a clear statistical picture "
        "of the pandemic's global outcome. The transition to an endemic state is now evident, but the data leaves a clear "
        "record of the lessons learned regarding spread, mortality, and the efficacy of modern medicine.")

    # Final Save
    output_path = "report/COVID19_Detailed_Corporate_Report.docx"
    doc.save(output_path)
    print(f"Corporate report saved to {output_path}")

if __name__ == "__main__":
    create_corporate_report()
